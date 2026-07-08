import streamlit as st
from google import genai  # Library BARU resmi dari Google
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
from datetime import datetime

# --- 1. PAGE CONFIGURATION & LAYOUT ---
st.set_page_config(page_title="Digibyaas - AI Modul Ajar", page_icon="🧬", layout="wide")

st.title("🧬 Digibyaas: Generator Modul Ajar Terintegrasi (Kelas 1-12)")
st.caption("© 2026 Digibyaas. Platform Penyelarasan Kurikulum Merdeka - Semua Jenjang SD, SMP, & SMA")

# --- 2. SIDEBAR SECURE AUTHENTICATION ---
with st.sidebar:
    st.header("🔑 Autentikasi API")
    st.markdown("Pastikan Anda menggunakan API Key baru yang berstatus **Active** dari Google AI Studio.")
    
    api_key_input = st.text_input(
        "Masukkan Gemini API Key Anda:", 
        type="password", 
        key="secure_gemini_api_key_pool"
    )
    
    st.markdown("---")
    st.header("🔬 Landasan Pedagogis")
    st.markdown("""
    Aplikasi ini merancang Modul Ajar otomatis dengan pilar:
    1. **Dual-Aspect Reflection**
    2. **Trilogi JMM** (Joyful, Mindful, Meaningful)
    """)

# --- 3. STATIC DATA DICTIONARIES ---
DEFAULT_SEKOLAH = "Sekolah Dasar & Menengah Nasional"
DEFAULT_KASEK = "Dr. H. Ahmad Sunarya, M.Pd."
DEFAULT_NIP_KASEK = "197508122000031002"
DEFAULT_GURU = "Budi Setiawan, S.Pd."
DEFAULT_NIP_GURU = "198810252015041003"

DATA_TINGKATAN = {
    "Kelas I (1) - SD [Fase A]": {"jenjang": "SD", "fase": "Fase A"},
    "Kelas II (2) - SD [Fase A]": {"jenjang": "SD", "fase": "Fase A"},
    "Kelas III (3) - SD [Fase B]": {"jenjang": "SD", "fase": "Fase B"},
    "Kelas IV (4) - SD [Fase B]": {"jenjang": "SD", "fase": "Fase B"},
    "Kelas V (5) - SD [Fase C]": {"jenjang": "SD", "fase": "Fase C"},
    "Kelas VI (6) - SD [Fase C]": {"jenjang": "SD", "fase": "Fase C"},
    "Kelas VII (7) - SMP [Fase D]": {"jenjang": "SMP", "fase": "Fase D"},
    "Kelas VIII (8) - SMP [Fase D]": {"jenjang": "SMP", "fase": "Fase D"},
    "Kelas IX (9) - SMP [Fase D]": {"jenjang": "SMP", "fase": "Fase D"},
    "Kelas X (10) - SMA [Fase E]": {"jenjang": "SMA", "fase": "Fase E"},
    "Kelas XI (11) - SMA [Fase F]": {"jenjang": "SMA", "fase": "Fase F"},
    "Kelas XII (12) - SMA [Fase F]": {"jenjang": "SMA", "fase": "Fase F"},
}

# --- 4. FORM INTERFACE GENERATION ---
st.subheader("📋 Identitas Sekolah & Kurikulum")
col1, col2 = st.columns(2)
with col1:
    nama_sekolah = st.text_input("Nama Sekolah", value=DEFAULT_SEKOLAH)
    pilihan_tingkat = st.selectbox("Pilih Kelas / Tingkatan", list(DATA_TINGKATAN.keys()))
    mata_pelajaran = st.text_input("Mata Pelajaran", placeholder="Contoh: Bahasa Indonesia / Matematika")
    materi_pokok = st.text_input("Materi Pokok", placeholder="Contoh: Penjumlahan Bilangan")
with col2:
    elemen = st.text_input("Elemen Pembelajaran", placeholder="Contoh: Aljabar / Bilangan")
    jumlah_jp = st.number_input("Jumlah Jam Pelajaran (JP)", min_value=1, max_value=10, value=2, step=1)
    jumlah_pertemuan = st.number_input("Untuk Berapa Pertemuan?", min_value=1, max_value=5, value=1, step=1)

info_tingkat = DATA_TINGKATAN[pilihan_tingkat]
jenjang_pendidikan = info_tingkat["jenjang"]
fase_kurikulum = info_tingkat["fase"]

st.subheader("🎯 Capaian & Tujuan")
cp = st.text_area("Capaian Pembelajaran (CP)")
atp = st.text_area("Tujuan Pembelajaran (TP)")

# --- FORM STRATEGI, MODEL, METODE, MEDIA & PENILAIAN ---
st.subheader("🛠️ Strategi & Asesmen Pembelajaran")
col5, col6 = st.columns(2)
with col5:
    model_pembelajaran = st.text_input("Model Pembelajaran", placeholder="Contoh: Problem Based Learning (PBL)")
    metode_pembelajaran = st.text_input("Metode Pembelajaran", placeholder="Contoh: Diskusi kelompok, Tanya Jawab")
with col6:
    bahan_ajar = st.text_area("Bahan Ajar & Media", placeholder="Contoh: LKPD Kelompok, Laptop, Proyektor")
    penilaian_asesmen = st.text_area("Jenis & Bentuk Penilaian", placeholder="Contoh: Formatif (Observasi sikap & LKPD)")

st.subheader("✍️ Otentikasi & Tanda Tangan")
col3, col4 = st.columns(2)
with col3:
    nama_kasek = st.text_input("Nama Kepala Sekolah", value=DEFAULT_KASEK)
    nip_kasek = st.text_input("NIP Kepala Sekolah", value=DEFAULT_NIP_KASEK)
with col4:
    nama_guru = st.text_input("Nama Guru Mata Pelajaran", value=DEFAULT_GURU)
    nip_guru = st.text_input("NIP Guru Mata Pelajaran", value=DEFAULT_GURU)


# --- FUNGSI PEMBUAT WORD (BERSIH BINTANG + TANDA TANGAN BERSEBELAHAN) ---
def buat_file_word(teks_rpp, nama_kasek, nip_kasek, nama_guru, nip_guru, tanggal, nama_sekolah):
    doc = Document()
    doc.add_heading('RENCANA PELAKSANAAN PEMBELAJARAN (RPP) / MODUL AJAR', level=1)
    
    isi_utama = []
    for baris in teks_rpp.split('\n'):
        if any(keyword in baris.lower() for keyword in ["kepala sekolah", "guru mata pelajaran", "mengetahui"]):
            continue
        if nip_kasek in baris or nip_guru in baris:
            continue
        isi_utama.append(baris)

    # Cetak isi konten utama
    for baris in isi_utama:
        p = doc.add_paragraph()
        bagian = baris.split('**')
        for i, teks in enumerate(bagian):
            if i % 2 == 1:
                p.add_run(teks).bold = True
            else:
                teks_whitespace = teks.replace('*', '')
                p.add_run(teks_whitespace)
                
    doc.add_paragraph() # Jarak sebelum tabel tanda tangan
    
    # Membuat tabel untuk tanda tangan bersebelahan (1 baris, 2 kolom)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Hilangkan border tabel agar bersih seperti kertas polos
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(tblBorders)
    
    # Kolom Kiri: Kepala Sekolah
    cell_kiri = table.rows[0].cells[0]
    p_kiri = cell_kiri.paragraphs[0]
    p_kiri.add_run(f"Mengetahui,\nKepala {nama_sekolah}\n\n\n\n").bold = False
    p_kiri.add_run(f"{nama_kasek}\n").bold = True
    p_kiri.add_run(f"NIP. {nip_kasek}")
    
    # Kolom Kanan: Guru Mata Pelajaran
    cell_kanan = table.rows[0].cells[1]
    p_kanan = cell_kanan.paragraphs[0]
    p_kanan.add_run(f", {tanggal}\nGuru Mata Pelajaran\n\n\n\n").bold = False
    p_kanan.add_run(f"{nama_guru}\n").bold = True
    p_kanan.add_run(f"NIP. {nip_guru}")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 5. EXECUTION ENGINE ---
if st.button("Susun RPP & Siapkan File Word"):
    if not api_key_input:
        st.error("🚨 Akses Ditolak! Masukkan Gemini API Key yang aktif pada input di sidebar kiri!")
    elif not cp or not atp or not mata_pelajaran or not materi_pokok:
        st.error("Mohon lengkapi data Kurikulum (Mapel, Materi Pokok, CP, dan TP)!")
    else:
        with st.spinner(f"Gemini AI sedang memproses data untuk {pilihan_tingkat}..."):
            try:
                # Inisialisasi Client google-genai
                fresh_client = genai.Client(api_key=api_key_input.strip())
                
                tanggal_sekarang = datetime.now().strftime("%d %B %Y")
                
                # Instruksi ketat struktur dengan optimasi "Tujuan Pembelajaran berbobot"
                prompt_instruksi = f"""
                Anda adalah pakar Kurikulum Merdeka tingkat {jenjang_pendidikan} ({fase_kurikulum}).
                Buatlah Modul Ajar/RPP yang praktis dan langsung pakai untuk kelas {pilihan_tingkat} di Sekolah {nama_sekolah}.
                
                INFORMASI UTAMA:
                - Nama Instansi/Sekolah: {nama_sekolah}
                - Mata Pelajaran: {mata_pelajaran}
                - Materi Pokok: {materi_pokok}
                - Elemen: {elemen}
                - Alokasi Waktu: {jumlah_jp} JP ({jumlah_pertemuan} Pertemuan)
                - Capaian Pembelajaran (CP): {cp}
                - Input Kompetensi: {atp}
                
                STRATEGI & ASESMEN:
                - Model Pembelajaran: {model_pembelajaran if model_pembelajaran else 'Sesuaikan yang relevan'}
                - Metode Pembelajaran: {metode_pembelajaran if metode_pembelajaran else 'Sesuaikan yang relevan'}
                - Bahan Ajar & Media: {bahan_ajar if bahan_ajar else 'Sesuaikan yang relevan'}
                - Penilaian / Asesmen: {penilaian_asesmen if penilaian_asesmen else 'Sesuaikan yang relevan'}

                ATURAN STRUKTUR MODUL AJAR (WAJIB DIPATUHI SECARA KETAT):
                1. Pada bagian "**IDENTITAS MODUL**", HILANGKAN dan JANGAN CETAK baris informasi berikut:
                   - Nama Penulis
                   - Jenjang Sekolah
                   - Tahun Pelajaran / Tahun Ajaran
                   Cukup tampilkan komponen: Instansi, Kelas, Mata Pelajaran, Materi Pokok, Elemen, dan Alokasi Waktu.
                2. Buatlah bagian bernama "**Tujuan Pembelajaran**". Di dalam bagian ini, jabarkan input kompetensi menjadi poin-poin tujuan yang operasional, detail, terukur, dan mendalam (gaya penulisan Indikator Ketercapaian/IKTP), TETAPI label judulnya harus TETAP ditulis "**Tujuan Pembelajaran**" (Jangan sebut istilah IKTP atau Indikator).
                3. JANGAN MEMBUAT atau menampilkan bagian-bagian berikut: Target Peserta Didik, Pemahaman Bermakna, dan Pertanyaan Esensial/Pemantik secara terpisah (HILANGKAN TOTAL).
                4. Struktur Kegiatan Pembelajaran hanya boleh terdiri dari:
                   - **Kegiatan Awal**: Masukkan Pertanyaan Pemantik langsung melebur ke dalam aktivitas awal ini bersama apersepsi dan motivasi (jangan buat sub-judul terpisah).
                   - **Kegiatan Inti**: Langkah aktivitas belajar siswa. Sematkan penanda Trilogi JMM [Joyful], [Mindful], atau [Meaningful] pada langkah yang sesuai.
                   - **Kegiatan Akhir**: Integrasikan bagian Refleksi Materi (untuk mengukur esensi materi pelajaran siswa) langsung melebur di sini sebelum kelas selesai.
                   - **Penutup**: Masukkan Refleksi Pembelajaran / Refleksi Proses (perasaan siswa dan evaluasi cara belajar mereka) langsung melebur ke dalam aktivitas penutup ini.
                5. Jangan menuliskan teks lembar pengesahan tanda tangan di akhir jawaban teks Anda, karena sistem kode Python akan mencetaknya secara otomatis secara berdampingan.
                """
                
                # Memanggil model gemini-2.5-flash
                response = fresh_client.models.generate_content(
                    model='gemini-2.5-flash', 
                    contents=prompt_instruksi
                )
                
                st.success("✨ RPP Resmi Selesai Dibuat!")
                st.markdown("---")
                st.markdown(response.text)
                
                # Proses convert teks ke Word dengan struktur tanda tangan tabel berdampingan yang bersih
                file_word = buat_file_word(
                    teks_rpp=response.text,
                    nama_kasek=nama_kasek,
                    nip_kasek=nip_kasek,
                    nama_guru=nama_guru,
                    nip_guru=nip_guru,
                    tanggal=tanggal_sekarang,
                    nama_sekolah=nama_sekolah
                )
                st.download_button(
                    label="📥 Download RPP (.docx)",
                    data=file_word,
                    file_name=f"RPP_{mata_pelajaran}_{pilihan_tingkat.split(' ')[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                error_msg = str(e)
                if "403" in error_msg or "API_KEY_INVALID" in error_msg:
                    st.error("🚨 Error 403: Kunci API salah atau tidak valid. Silakan buat API Key baru di Google AI Studio!")
                else:
                    st.error(f"Terjadi kendala pada sistem: {error_msg}")
